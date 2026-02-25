#!/usr/bin/env python3
"""
WSU Design Standards Compliance Review — Local Report Generator
Reads review-data.json and generates all deliverable files with validated numbers.

Usage: python generate_reports.py <path-to-review-data.json>

Outputs (written to same directory as input JSON):
  - report.docx    Word report with WSU crimson styling
  - report.xlsx    Excel workbook (5 sheets, Summary prints to 1 landscape page)
  - checklist.txt  Plain-text checklist
  - findings.txt   Plain-text findings log
  - notes.txt      Plain-text methodology notes
  - COMPLETE       Sentinel (only if validation passes)
  - FAILED         Error file (only if validation fails)
"""

import json, os, sys, traceback
from datetime import datetime

# ---------------------------------------------------------------------------
# Lazy imports — installed by pipeline if needed
# ---------------------------------------------------------------------------
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.worksheet.page import PageMargins
except ImportError:
    sys.exit("ERROR: openpyxl not installed. Run: pip install openpyxl")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("ERROR: python-docx not installed. Run: pip install python-docx")

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
WSU_CRIMSON   = "981E32"
WSU_GRAY      = "5F6B6D"
DARK_GRAY     = "374151"
WHITE_HEX     = "FFFFFF"
LIGHT_RED     = "FEE2E2"
LIGHT_ORANGE  = "FEF3C7"
LIGHT_YELLOW  = "FEF9C3"
LIGHT_GREEN   = "DCFCE7"
LIGHT_BLUE    = "EFF6FF"
ALT_ROW       = "F9FAFB"

SEVERITY_ORDER = {"Critical": 0, "Major": 1, "Minor": 2}
SEVERITY_FILL = {
    "Critical": LIGHT_RED,
    "Major":    LIGHT_ORANGE,
    "Minor":    LIGHT_YELLOW,
}

DIVISION_NAMES = {
    "02": "Existing Conditions", "03": "Concrete", "04": "Masonry",
    "05": "Metals", "07": "Thermal & Moisture Protection", "08": "Openings",
    "09": "Finishes", "10": "Specialties", "11": "Equipment",
    "12": "Furnishings", "13": "Special Construction", "14": "Conveying Equipment",
    "21": "Fire Suppression", "22": "Plumbing", "23": "HVAC / Mechanical",
    "25": "Integrated Automation", "26": "Electrical", "27": "Communications",
    "28": "Electronic Safety & Security", "31": "Earthwork",
    "32": "Exterior Improvements", "33": "Utilities", "40": "Process Integration",
}

THIN_BORDER = Border(
    left=Side(style="thin"), right=Side(style="thin"),
    top=Side(style="thin"), bottom=Side(style="thin"),
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def sev_sort_key(f):
    return (SEVERITY_ORDER.get(f.get("severity", ""), 99), f.get("division", "99"))

def pct(num, denom):
    return f"{num / denom * 100:.1f}%" if denom > 0 else "0.0%"

def cell_fill(hex_color):
    return PatternFill(start_color=hex_color, end_color=hex_color, fill_type="solid")

def set_cell_shading(cell, hex_color):
    """Set cell background via XML for both openpyxl and docx cells."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)

# ---------------------------------------------------------------------------
# Validation
# ---------------------------------------------------------------------------
def validate(data, findings):
    """Cross-check all numbers. Returns list of error strings (empty = pass)."""
    errors = []

    # Count findings by severity
    sev_counts = {"Critical": 0, "Major": 0, "Minor": 0}
    for i, f in enumerate(findings):
        s = f.get("severity", "")
        if s in sev_counts:
            sev_counts[s] += 1
        elif s:
            errors.append(f"Finding {f.get('id','?')}: unknown severity '{s}'")
        else:
            errors.append(f"Finding {f.get('id','?')}: missing severity")

    # Count non-compliant from discipline summaries
    disc_nc = 0
    disc_total_req = 0
    for d in data.get("disciplines", []):
        s = d.get("summary", {})
        row_d = s.get("deviations", 0)
        row_o = s.get("omissions", 0)
        row_x = s.get("concerns", 0)
        row_c = s.get("compliant", 0)
        row_total = s.get("total_requirements", 0)
        disc_nc += row_d + row_o + row_x
        disc_total_req += row_total
        # Per-row balance
        if row_c + row_d + row_o + row_x != row_total:
            errors.append(
                f"Discipline {d['key']}: C({row_c})+D({row_d})+O({row_o})+X({row_x}) "
                f"= {row_c+row_d+row_o+row_x} != Total({row_total})"
            )

    # Findings must have at least 1 entry
    if len(findings) == 0:
        errors.append("No findings in findings array")

    # Findings cannot exceed discipline non-compliant total
    if len(findings) > disc_nc and disc_nc > 0:
        errors.append(
            f"Findings count ({len(findings)}) exceeds discipline "
            f"non-compliant total ({disc_nc})"
        )

    # Requirements list count should match disc_total_req (if provided)
    req_count = len(data.get("requirements", []))
    if req_count > 0 and req_count != disc_total_req:
        errors.append(f"Requirements list ({req_count}) != discipline total ({disc_total_req})")

    return errors

# ---------------------------------------------------------------------------
# Text file generators
# ---------------------------------------------------------------------------
def generate_findings_txt(findings, output_dir):
    lines = ["WSU DESIGN STANDARDS COMPLIANCE REVIEW — FINDINGS LOG", "=" * 60, ""]
    lines.append(f"Total Findings: {len(findings)}")
    sev = {"Critical": 0, "Major": 0, "Minor": 0}
    for f in findings:
        sev[f.get("severity", "Minor")] = sev.get(f.get("severity", "Minor"), 0) + 1
    lines.append(f"Critical: {sev['Critical']}  |  Major: {sev['Major']}  |  Minor: {sev['Minor']}")
    lines.append("")
    lines.append("-" * 60)
    for f in findings:
        lines.append("")
        lines.append(f"{f['id']}: {f.get('title', 'Untitled')}")
        lines.append(f"  Division:    {f.get('division', '??')} — {DIVISION_NAMES.get(f.get('division',''), 'Unknown')}")
        lines.append(f"  CSI Code:    {f.get('csi_code', '')}")
        lines.append(f"  Severity:    {f.get('severity', '')}")
        lines.append(f"  Status:      [{f.get('status', '?')}]")
        lines.append(f"  PDF Ref:     {f.get('pdf_reference', '')}")
        lines.append(f"  Standard:    {f.get('standard_reference', '')}")
        lines.append(f"  Issue:       {f.get('issue', '')}")
        lines.append(f"  Action:      {f.get('required_action', '')}")
        lines.append("-" * 60)
    with open(os.path.join(output_dir, "findings.txt"), "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))

def generate_checklist_txt(data, findings, output_dir):
    lines = ["WSU DESIGN STANDARDS COMPLIANCE REVIEW — CHECKLIST", "=" * 60, ""]
    proj = data.get("project", {})
    lines.append(f"Project: {proj.get('name', '')}")
    lines.append(f"Phase:   {proj.get('phase', '')}")
    lines.append(f"Date:    {proj.get('reviewDate', '')}")
    lines.append("")
    # Build finding ref lookup
    fref = {}
    for f in findings:
        key = (f.get("division", ""), f.get("csi_code", ""), f.get("status", ""))
        if key not in fref:
            fref[key] = f["id"]
    lines.append(f"{'#':<5} {'Div':<5} {'CSI Code':<14} {'Requirement':<50} {'Status':<8} {'Finding':<8} {'Sheet':<12} Notes")
    lines.append("-" * 140)
    reqs = data.get("requirements", [])
    for i, r in enumerate(reqs, 1):
        status = r.get("status", "C")
        ref = r.get("finding_ref", "") or ""
        lines.append(
            f"{i:<5} {r.get('division',''):<5} {r.get('csi_code',''):<14} "
            f"{(r.get('requirement','') or '')[:50]:<50} [{status}]{'':>4} "
            f"{ref:<8} {(r.get('drawing_sheet','') or ''):<12} {r.get('notes','') or ''}"
        )
    lines.append("-" * 140)
    with open(os.path.join(output_dir, "checklist.txt"), "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))

def generate_notes_txt(data, output_dir):
    narr = data.get("narratives", {})
    lines = ["WSU DESIGN STANDARDS COMPLIANCE REVIEW — NOTES", "=" * 60, ""]
    lines.append("METHODOLOGY")
    lines.append("-" * 40)
    lines.append(narr.get("methodology", "Automated review against WSU Design Standards."))
    lines.append("")
    lines.append("DRAWING INVENTORY")
    lines.append("-" * 40)
    lines.append(narr.get("drawing_inventory", "See PDF drawing set."))
    lines.append("")
    lines.append("STANDARDS APPLICABILITY")
    lines.append("-" * 40)
    lines.append(narr.get("applicability_notes", "All selected divisions reviewed."))
    lines.append("")
    lines.append("REVIEW LIMITATIONS")
    lines.append("-" * 40)
    lines.append(narr.get("limitations", "Review based on drawing set only; specifications not available."))
    with open(os.path.join(output_dir, "notes.txt"), "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))

# ---------------------------------------------------------------------------
# Excel generator
# ---------------------------------------------------------------------------
def generate_xlsx(data, findings, output_dir):
    proj = data.get("project", {})
    disciplines = data.get("disciplines", [])
    requirements = data.get("requirements", [])
    variances = data.get("variances", [])
    narr = data.get("narratives", {})

    wb = openpyxl.Workbook()

    hdr_font = Font(name="Calibri", size=9, bold=True, color=WHITE_HEX)
    hdr_fill = cell_fill(DARK_GRAY)
    body_font = Font(name="Calibri", size=9)
    bold_font = Font(name="Calibri", size=9, bold=True)
    section_font = Font(name="Calibri", size=11, bold=True, color=WSU_CRIMSON)
    title_font = Font(name="Calibri", size=14, bold=True, color=WSU_CRIMSON)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_wrap = Alignment(horizontal="left", vertical="top", wrap_text=True)
    total_fill = cell_fill(LIGHT_BLUE)

    def style_header(ws, row, col_count):
        for c in range(1, col_count + 1):
            cell = ws.cell(row=row, column=c)
            cell.font = hdr_font
            cell.fill = hdr_fill
            cell.alignment = center
            cell.border = THIN_BORDER

    def style_body_cell(cell, align="left"):
        cell.font = body_font
        cell.border = THIN_BORDER
        if align == "center":
            cell.alignment = center
        elif align == "right":
            cell.alignment = Alignment(horizontal="right", vertical="center")
        else:
            cell.alignment = left_wrap

    # ── Sheet 1: Summary ──────────────────────────────────────────────────
    ws = wb.active
    ws.title = "Summary"
    ws.page_setup.orientation = "landscape"
    ws.page_setup.paperSize = ws.PAPERSIZE_LETTER
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_margins = PageMargins(left=0.5, right=0.5, top=0.5, bottom=0.5, header=0.3, footer=0.3)
    ws.oddFooter.center.text = "WSU Facilities Services -- Design Standards Compliance Review"
    ws.oddFooter.right.text = "Page &P"

    col_widths = [3, 18, 14, 14, 14, 14, 14, 14, 14]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # Header
    ws["B1"] = "WASHINGTON STATE UNIVERSITY"
    ws["B1"].font = title_font
    ws["B2"] = "Design Standards Compliance Review"
    ws["B2"].font = Font(name="Calibri", size=11, bold=True)

    # Project info
    ws["B3"] = "Project:"
    ws["B3"].font = bold_font
    ws["C3"] = proj.get("name", "")
    ws["D3"] = "Phase:"
    ws["D3"].font = bold_font
    ws["E3"] = proj.get("phase", "")
    ws["B4"] = "Review Date:"
    ws["B4"].font = bold_font
    ws["C4"] = proj.get("reviewDate", "")
    ws["D4"] = "Construction Type:"
    ws["D4"].font = bold_font
    ws["E4"] = proj.get("constructionType", "")

    # Severity counts
    sev = {"Critical": 0, "Major": 0, "Minor": 0}
    for f in findings:
        s = f.get("severity", "Minor")
        if s in sev:
            sev[s] += 1
    total_nc = sum(sev.values())

    # Total requirements
    total_req = sum(d.get("summary", {}).get("total_requirements", 0) for d in disciplines)
    total_compliant = sum(d.get("summary", {}).get("compliant", 0) for d in disciplines)
    compliance_pct = total_compliant / total_req * 100 if total_req > 0 else 0

    # Executive summary
    ws["B6"] = "EXECUTIVE SUMMARY"
    ws["B6"].font = section_font
    summary_text = narr.get("executive_summary", "")
    if not summary_text:
        summary_text = (
            f"This review evaluated {total_req} requirements across {len(disciplines)} disciplines. "
            f"{compliance_pct:.1f}% of requirements are compliant. "
            f"{sev['Critical']} critical findings require immediate attention."
        )
    ws.merge_cells("B7:I8")
    ws["B7"] = summary_text
    ws["B7"].font = body_font
    ws["B7"].alignment = left_wrap

    # Findings by Severity
    row = 10
    ws.cell(row=row, column=2, value="FINDINGS BY SEVERITY").font = section_font
    row += 1
    for c, h in enumerate(["Severity", "Count", "% of Findings"], 2):
        ws.cell(row=row, column=c, value=h)
    style_header(ws, row, 4)

    sev_rows = [
        ("Critical", sev["Critical"], LIGHT_RED),
        ("Major", sev["Major"], LIGHT_ORANGE),
        ("Minor", sev["Minor"], LIGHT_YELLOW),
    ]
    for label, count, fill_color in sev_rows:
        row += 1
        ws.cell(row=row, column=2, value=label)
        ws.cell(row=row, column=3, value=count)
        ws.cell(row=row, column=4, value=pct(count, total_nc))
        for c in range(2, 5):
            cell = ws.cell(row=row, column=c)
            cell.fill = cell_fill(fill_color)
            style_body_cell(cell, "center" if c > 2 else "left")

    row += 1
    ws.cell(row=row, column=2, value="Total Non-Compliant").font = bold_font
    ws.cell(row=row, column=3, value=total_nc).font = bold_font
    ws.cell(row=row, column=4, value="100%").font = bold_font
    for c in range(2, 5):
        ws.cell(row=row, column=c).border = THIN_BORDER

    # Findings by Discipline
    row += 2
    ws.cell(row=row, column=2, value="FINDINGS BY DISCIPLINE").font = section_font
    row += 1
    disc_headers = ["Division", "Description", "Total Req.", "Compliant",
                    "Deviations", "Omissions", "Concerns", "Compliance %"]
    for c, h in enumerate(disc_headers, 2):
        ws.cell(row=row, column=c, value=h)
    style_header(ws, row, 9)

    # Build per-division rows from disciplines
    div_rows = []
    for disc in disciplines:
        s = disc.get("summary", {})
        for div in sorted(disc.get("divisions_reviewed", [])):
            # Count per-division from requirements if available
            div_reqs = [r for r in requirements if r.get("division") == div]
            if div_reqs:
                d_total = len(div_reqs)
                d_c = sum(1 for r in div_reqs if r.get("status") == "C")
                d_d = sum(1 for r in div_reqs if r.get("status") == "D")
                d_o = sum(1 for r in div_reqs if r.get("status") == "O")
                d_x = sum(1 for r in div_reqs if r.get("status") == "X")
            else:
                # Fallback: use discipline summary (not per-division)
                d_total = s.get("total_requirements", 0)
                d_c = s.get("compliant", 0)
                d_d = s.get("deviations", 0)
                d_o = s.get("omissions", 0)
                d_x = s.get("concerns", 0)
            div_rows.append({
                "div": div,
                "desc": DIVISION_NAMES.get(div, "Unknown"),
                "total": d_total, "c": d_c, "d": d_d, "o": d_o, "x": d_x,
            })

    # Deduplicate and sort
    seen = set()
    unique_rows = []
    for dr in sorted(div_rows, key=lambda x: x["div"]):
        if dr["div"] not in seen:
            seen.add(dr["div"])
            unique_rows.append(dr)

    totals = {"total": 0, "c": 0, "d": 0, "o": 0, "x": 0}
    for dr in unique_rows:
        row += 1
        vals = [dr["div"], dr["desc"], dr["total"], dr["c"], dr["d"], dr["o"], dr["x"],
                pct(dr["c"], dr["total"])]
        for c, v in enumerate(vals, 2):
            cell = ws.cell(row=row, column=c, value=v)
            style_body_cell(cell, "center" if c > 3 else "left")
        for k in ["total", "c", "d", "o", "x"]:
            totals[k] += dr[k]

    # Total row
    row += 1
    total_vals = ["TOTAL", "", totals["total"], totals["c"], totals["d"],
                  totals["o"], totals["x"], pct(totals["c"], totals["total"])]
    for c, v in enumerate(total_vals, 2):
        cell = ws.cell(row=row, column=c, value=v)
        cell.font = bold_font
        cell.fill = total_fill
        cell.border = THIN_BORDER

    # Top Critical Findings
    row += 2
    ws.cell(row=row, column=2, value="TOP CRITICAL FINDINGS").font = section_font
    row += 1
    criticals = [f for f in findings if f.get("severity") == "Critical"]
    if criticals:
        crit_headers = ["#", "Finding", "Division", "Description", "Required Action"]
        for c, h in enumerate(crit_headers, 2):
            ws.cell(row=row, column=c, value=h)
        style_header(ws, row, 6)
        for i, cf in enumerate(criticals[:6], 1):
            row += 1
            vals = [i, cf["id"], cf.get("division", ""),
                    (cf.get("title", "") or cf.get("issue", ""))[:80],
                    (cf.get("required_action", ""))[:80]]
            for c, v in enumerate(vals, 2):
                cell = ws.cell(row=row, column=c, value=v)
                style_body_cell(cell)
                cell.fill = cell_fill(LIGHT_RED)
    else:
        ws.cell(row=row, column=2, value="No critical findings identified.")

    # Priority Actions
    row += 2
    ws.cell(row=row, column=2, value="PRIORITY ACTIONS").font = section_font
    row += 1
    prio_headers = ["Priority", "Timeline", "Finding Count"]
    for c, h in enumerate(prio_headers, 2):
        ws.cell(row=row, column=c, value=h)
    style_header(ws, row, 4)
    prio_data = [
        ("Immediate (Critical)", "Before next submission", sev["Critical"]),
        ("CD Phase (Major)", "Resolve during CD", sev["Major"]),
        ("100% CD (Minor)", "Address by final CD", sev["Minor"]),
    ]
    for label, timeline, count in prio_data:
        row += 1
        for c, v in enumerate([label, timeline, count], 2):
            cell = ws.cell(row=row, column=c, value=v)
            style_body_cell(cell, "center" if c == 4 else "left")

    # Set print area
    ws.print_area = f"A1:I{row}"

    # ── Sheet 2: Findings ─────────────────────────────────────────────────
    ws2 = wb.create_sheet("Findings")
    ws2.freeze_panes = "A2"
    col_widths2 = [6, 10, 12, 10, 8, 12, 16, 40, 40]
    for i, w in enumerate(col_widths2, 1):
        ws2.column_dimensions[get_column_letter(i)].width = w

    headers2 = ["F-#", "Division", "CSI Code", "Severity", "Status",
                "PDF Sheet", "Standard Ref", "Issue", "Required Action"]
    for c, h in enumerate(headers2, 1):
        ws2.cell(row=1, column=c, value=h)
    style_header(ws2, 1, 9)
    ws2.auto_filter.ref = f"A1:I{len(findings) + 1}"

    for i, f in enumerate(findings, 2):
        vals = [f["id"], f.get("division", ""), f.get("csi_code", ""),
                f.get("severity", ""), f"[{f.get('status', '?')}]",
                f.get("pdf_reference", ""), f.get("standard_reference", ""),
                f.get("issue", ""), f.get("required_action", "")]
        fill = cell_fill(SEVERITY_FILL.get(f.get("severity", ""), LIGHT_YELLOW))
        for c, v in enumerate(vals, 1):
            cell = ws2.cell(row=i, column=c, value=v)
            style_body_cell(cell)
            cell.fill = fill

    # ── Sheet 3: Checklist ────────────────────────────────────────────────
    ws3 = wb.create_sheet("Checklist")
    ws3.freeze_panes = "A2"
    col_widths3 = [5, 10, 12, 50, 8, 8, 12, 30]
    for i, w in enumerate(col_widths3, 1):
        ws3.column_dimensions[get_column_letter(i)].width = w

    headers3 = ["#", "Division", "CSI Code", "Requirement", "Status",
                "Finding Ref", "Drawing Sheet", "Notes"]
    for c, h in enumerate(headers3, 1):
        ws3.cell(row=1, column=c, value=h)
    style_header(ws3, 1, 8)
    ws3.auto_filter.ref = f"A1:H{len(requirements) + 1}"

    status_fill = {"C": LIGHT_GREEN, "D": LIGHT_RED, "O": LIGHT_ORANGE, "X": LIGHT_YELLOW}
    for i, r in enumerate(requirements, 2):
        st = r.get("status", "C")
        vals = [i - 1, r.get("division", ""), r.get("csi_code", ""),
                r.get("requirement", ""), f"[{st}]",
                r.get("finding_ref", "") or "", r.get("drawing_sheet", "") or "",
                r.get("notes", "") or ""]
        fill = cell_fill(status_fill.get(st, WHITE_HEX))
        for c, v in enumerate(vals, 1):
            cell = ws3.cell(row=i, column=c, value=v)
            style_body_cell(cell)
            cell.fill = fill

    # ── Sheet 4: Variances ────────────────────────────────────────────────
    ws4 = wb.create_sheet("Variances")
    ws4.freeze_panes = "A2"
    headers4 = ["V-#", "Finding Ref", "Division", "CSI Code",
                "Description", "Justification", "Approval Status"]
    for c, h in enumerate(headers4, 1):
        ws4.cell(row=1, column=c, value=h)
    style_header(ws4, 1, 7)

    if variances:
        for i, v in enumerate(variances, 2):
            vals = [v.get("id", f"V-{i-1:02d}"), v.get("finding_ref", ""),
                    v.get("division", ""), v.get("csi_code", ""),
                    v.get("description", ""), v.get("justification", ""),
                    v.get("approval_status", "Pending")]
            for c, val in enumerate(vals, 1):
                cell = ws4.cell(row=i, column=c, value=val)
                style_body_cell(cell)
    else:
        ws4.cell(row=2, column=2, value="No variance requests identified for this project.")

    # ── Sheet 5: Notes ────────────────────────────────────────────────────
    ws5 = wb.create_sheet("Notes")
    narr = data.get("narratives", {})
    ws5["A1"] = "REVIEW NOTES"
    ws5["A1"].font = Font(name="Calibri", size=14, bold=True)
    row = 3
    for section, key in [("Methodology", "methodology"),
                         ("Drawing Inventory", "drawing_inventory"),
                         ("Standards Applicability", "applicability_notes"),
                         ("Review Limitations", "limitations")]:
        ws5.cell(row=row, column=1, value=section).font = section_font
        row += 1
        ws5.merge_cells(f"A{row}:H{row}")
        ws5.cell(row=row, column=1, value=narr.get(key, "")).alignment = left_wrap
        row += 2

    wb.save(os.path.join(output_dir, "report.xlsx"))

# ---------------------------------------------------------------------------
# Word generator
# ---------------------------------------------------------------------------
def generate_docx(data, findings, output_dir):
    proj = data.get("project", {})
    disciplines = data.get("disciplines", [])
    requirements = data.get("requirements", [])
    variances = data.get("variances", [])
    narr = data.get("narratives", {})

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Styles
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(10)
    style_normal.font.color.rgb = RGBColor(0, 0, 0)

    for level, size, color in [
        ("Heading 1", 16, RGBColor(0x98, 0x1E, 0x32)),
        ("Heading 2", 13, RGBColor(0x37, 0x41, 0x51)),
        ("Heading 3", 11, RGBColor(0, 0, 0)),
    ]:
        s = doc.styles[level]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.color.rgb = color
        s.font.bold = True

    # Header/Footer
    header_para = section.header.paragraphs[0]
    header_run = header_para.add_run(f"WSU Design Standards Compliance Review - {proj.get('name', '')}")
    header_run.font.size = Pt(8)
    header_run.font.color.rgb = RGBColor(0x5F, 0x6B, 0x6D)
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer_para = section.footer.paragraphs[0]
    footer_run = footer_para.add_run("WSU Facilities Services - Confidential")
    footer_run.font.size = Pt(8)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_crimson_rule():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), WSU_CRIMSON)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def format_table(table):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell in table.rows[0].cells:
            set_cell_shading(cell, DARK_GRAY)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(9)
        for ri, row in enumerate(table.rows[1:], 1):
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
                        run.font.name = "Calibri"
                if ri % 2 == 0:
                    set_cell_shading(cell, ALT_ROW)

    # ── Cover Page ─────────────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("WASHINGTON STATE UNIVERSITY")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x98, 0x1E, 0x32)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Facilities Services")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x5F, 0x6B, 0x6D)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Design Standards Compliance Review")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    add_crimson_rule()

    # Project info table (no borders)
    info_fields = [
        ("Project:", proj.get("name", "")),
        ("Review Phase:", proj.get("phase", "")),
        ("Construction Type:", proj.get("constructionType", "")),
        ("Review Date:", proj.get("reviewDate", "")),
        ("Prepared by:", "WSU Facilities Services - Automated Compliance Review"),
    ]
    t = doc.add_table(rows=len(info_fields), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, val) in enumerate(info_fields):
        t.rows[i].cells[0].text = label
        t.rows[i].cells[1].text = val
        for run in t.rows[i].cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(12)
        for run in t.rows[i].cells[1].paragraphs[0].runs:
            run.font.size = Pt(12)

    doc.add_page_break()

    # ── Section 2: Executive Summary ───────────────────────────────────────
    doc.add_heading("Executive Summary", level=1)
    add_crimson_rule()
    sev = {"Critical": 0, "Major": 0, "Minor": 0}
    for f in findings:
        s = f.get("severity", "Minor")
        if s in sev:
            sev[s] += 1
    total_nc = sum(sev.values())
    total_req = sum(d.get("summary", {}).get("total_requirements", 0) for d in disciplines)
    total_compliant = sum(d.get("summary", {}).get("compliant", 0) for d in disciplines)

    exec_text = narr.get("executive_summary", "")
    if not exec_text:
        exec_text = (
            f"This review evaluated {total_req} requirements across {len(disciplines)} "
            f"disciplines. {total_compliant / total_req * 100:.1f}% of requirements are compliant. "
            f"{sev['Critical']} critical findings require immediate attention."
        )
    doc.add_paragraph(exec_text)
    doc.add_page_break()

    # ── Section 3: Scope & Methodology ────────────────────────────────────
    doc.add_heading("Review Scope & Methodology", level=1)
    add_crimson_rule()
    doc.add_paragraph(narr.get("methodology", "Automated review against current WSU Design Standards."))
    doc.add_page_break()

    # ── Section 4: Findings Summary ───────────────────────────────────────
    doc.add_heading("Findings Summary by Discipline", level=1)
    add_crimson_rule()

    # Severity table
    doc.add_heading("Findings by Severity", level=2)
    t = doc.add_table(rows=4, cols=3)
    t.style = "Table Grid"
    headers = ["Severity", "Count", "% of Findings"]
    for c, h in enumerate(headers):
        t.rows[0].cells[c].text = h
    sev_data = [
        ("Critical", sev["Critical"]),
        ("Major", sev["Major"]),
        ("Minor", sev["Minor"]),
    ]
    for i, (label, count) in enumerate(sev_data):
        t.rows[i + 1].cells[0].text = label
        t.rows[i + 1].cells[1].text = str(count)
        t.rows[i + 1].cells[2].text = pct(count, total_nc)
        fill_hex = SEVERITY_FILL.get(label, LIGHT_YELLOW)
        for cell in t.rows[i + 1].cells:
            set_cell_shading(cell, fill_hex)
    format_table(t)

    p = doc.add_paragraph()
    run = p.add_run(f"Total Non-Compliant: {total_nc}")
    run.font.bold = True

    # Discipline table
    doc.add_heading("Findings by Discipline", level=2)
    disc_headers = ["Division", "Description", "Total", "Compliant", "Dev.", "Omis.", "Concerns", "Comp. %"]
    # Collect unique divisions from requirements
    div_data = {}
    for r in requirements:
        div = r.get("division", "")
        if div not in div_data:
            div_data[div] = {"total": 0, "c": 0, "d": 0, "o": 0, "x": 0}
        div_data[div]["total"] += 1
        st = r.get("status", "C")
        if st == "C": div_data[div]["c"] += 1
        elif st == "D": div_data[div]["d"] += 1
        elif st == "O": div_data[div]["o"] += 1
        elif st == "X": div_data[div]["x"] += 1

    sorted_divs = sorted(div_data.keys())
    t = doc.add_table(rows=len(sorted_divs) + 2, cols=8)
    t.style = "Table Grid"
    for c, h in enumerate(disc_headers):
        t.rows[0].cells[c].text = h
    for i, div in enumerate(sorted_divs, 1):
        d = div_data[div]
        vals = [div, DIVISION_NAMES.get(div, ""), str(d["total"]), str(d["c"]),
                str(d["d"]), str(d["o"]), str(d["x"]), pct(d["c"], d["total"])]
        for c, v in enumerate(vals):
            t.rows[i].cells[c].text = v
    # Total row
    tr = len(sorted_divs) + 1
    t_vals = ["TOTAL", "", str(sum(d["total"] for d in div_data.values())),
              str(sum(d["c"] for d in div_data.values())),
              str(sum(d["d"] for d in div_data.values())),
              str(sum(d["o"] for d in div_data.values())),
              str(sum(d["x"] for d in div_data.values())),
              pct(sum(d["c"] for d in div_data.values()),
                  sum(d["total"] for d in div_data.values()))]
    for c, v in enumerate(t_vals):
        t.rows[tr].cells[c].text = v
        for run in t.rows[tr].cells[c].paragraphs[0].runs:
            run.font.bold = True
        set_cell_shading(t.rows[tr].cells[c], LIGHT_BLUE)
    format_table(t)
    doc.add_page_break()

    # ── Section 5: Top Critical Findings ──────────────────────────────────
    doc.add_heading("Top Critical Findings", level=1)
    add_crimson_rule()
    criticals = [f for f in findings if f.get("severity") == "Critical"]
    if criticals:
        for cf in criticals:
            doc.add_heading(f"{cf['id']}: {cf.get('title', '')}", level=3)
            meta = [
                ("Division", f"{cf.get('division', '')} - {DIVISION_NAMES.get(cf.get('division', ''), '')}"),
                ("CSI Code", cf.get("csi_code", "")),
                ("Severity", "Critical"),
                ("Status", f"[{cf.get('status', '?')}]"),
                ("PDF Reference", cf.get("pdf_reference", "")),
                ("Standard", cf.get("standard_reference", "")),
            ]
            for label, val in meta:
                p = doc.add_paragraph()
                run = p.add_run(f"{label}: ")
                run.font.bold = True
                run.font.size = Pt(9)
                run = p.add_run(val)
                run.font.size = Pt(9)
                p.paragraph_format.space_after = Pt(1)
            p = doc.add_paragraph()
            run = p.add_run("Issue: ")
            run.font.bold = True
            p.add_run(cf.get("issue", ""))
            p = doc.add_paragraph()
            run = p.add_run("Required Action: ")
            run.font.bold = True
            p.add_run(cf.get("required_action", ""))
    else:
        doc.add_paragraph("No critical findings identified.")
    doc.add_page_break()

    # ── Section 6: Detailed Findings by Discipline ────────────────────────
    doc.add_heading("Detailed Findings by Discipline", level=1)
    add_crimson_rule()

    # Group findings by discipline
    by_disc = {}
    for f in findings:
        disc = f.get("discipline", "unknown")
        by_disc.setdefault(disc, []).append(f)

    for disc_key in sorted(by_disc.keys()):
        disc_findings = by_disc[disc_key]
        # Get discipline name
        disc_name = disc_key.replace("-", " ").title()
        for d in disciplines:
            if d.get("key") == disc_key:
                disc_name = d.get("name", disc_name)
                break
        doc.add_heading(disc_name, level=2)
        for f in disc_findings:
            doc.add_heading(f"{f['id']}: {f.get('title', '')}", level=3)
            meta = [
                ("Division", f"{f.get('division', '')}"),
                ("CSI Code", f.get("csi_code", "")),
                ("Severity", f.get("severity", "")),
                ("Status", f"[{f.get('status', '?')}]"),
                ("PDF Reference", f.get("pdf_reference", "")),
                ("Standard", f.get("standard_reference", "")),
            ]
            for label, val in meta:
                p = doc.add_paragraph()
                run = p.add_run(f"{label}: ")
                run.font.bold = True
                run.font.size = Pt(9)
                run = p.add_run(val)
                run.font.size = Pt(9)
                p.paragraph_format.space_after = Pt(1)
            if f.get("issue"):
                p = doc.add_paragraph()
                run = p.add_run("Issue: ")
                run.font.bold = True
                p.add_run(f.get("issue", ""))
            if f.get("required_action"):
                p = doc.add_paragraph()
                run = p.add_run("Required Action: ")
                run.font.bold = True
                p.add_run(f.get("required_action", ""))
    doc.add_page_break()

    # ── Section 7: Variances ──────────────────────────────────────────────
    doc.add_heading("Items Requiring Formal WSU Variance Approval", level=1)
    add_crimson_rule()
    if variances:
        t = doc.add_table(rows=len(variances) + 1, cols=5)
        t.style = "Table Grid"
        for c, h in enumerate(["V-#", "Finding", "Division", "Description", "Status"]):
            t.rows[0].cells[c].text = h
        for i, v in enumerate(variances, 1):
            vals = [v.get("id", ""), v.get("finding_ref", ""), v.get("division", ""),
                    v.get("description", ""), v.get("approval_status", "Pending")]
            for c, val in enumerate(vals):
                t.rows[i].cells[c].text = val
        format_table(t)
    else:
        doc.add_paragraph("No variance requests identified for this project.")
    doc.add_page_break()

    # ── Section 8: Applicability Notes ────────────────────────────────────
    doc.add_heading("Standards Applicability Notes", level=1)
    add_crimson_rule()
    doc.add_paragraph(narr.get("applicability_notes", "All selected divisions were reviewed."))
    doc.add_page_break()

    # ── Section 9: Priority Action Matrix ─────────────────────────────────
    doc.add_heading("Priority Action Matrix", level=1)
    add_crimson_rule()

    for priority, sev_level in [("Immediate Action Required", "Critical"),
                                 ("Resolve During CD Phase", "Major"),
                                 ("Resolve by 100% CD Submission", "Minor")]:
        doc.add_heading(priority, level=2)
        pf = [f for f in findings if f.get("severity") == sev_level]
        if pf:
            t = doc.add_table(rows=len(pf) + 1, cols=4)
            t.style = "Table Grid"
            for c, h in enumerate(["Finding #", "Division", "Description", "Severity"]):
                t.rows[0].cells[c].text = h
            for i, f in enumerate(pf, 1):
                vals = [f["id"], f.get("division", ""),
                        (f.get("title") or f.get("issue") or "")[:60], sev_level]
                for c, v in enumerate(vals):
                    t.rows[i].cells[c].text = v
            format_table(t)
        else:
            doc.add_paragraph(f"No {sev_level.lower()} findings.")

    # Summary table
    doc.add_heading("Summary", level=2)
    t = doc.add_table(rows=4, cols=3)
    t.style = "Table Grid"
    for c, h in enumerate(["Priority Level", "Count", "Target Resolution"]):
        t.rows[0].cells[c].text = h
    prio_sum = [
        ("Immediate", str(sev["Critical"]), "Before next submission"),
        ("CD Phase", str(sev["Major"]), "During CD development"),
        ("100% CD", str(sev["Minor"]), "By final CD submission"),
    ]
    for i, (lbl, cnt, tgt) in enumerate(prio_sum, 1):
        t.rows[i].cells[0].text = lbl
        t.rows[i].cells[1].text = cnt
        t.rows[i].cells[2].text = tgt
    format_table(t)
    doc.add_page_break()

    # ── Section 10: Next Steps ────────────────────────────────────────────
    doc.add_heading("Next Steps", level=1)
    add_crimson_rule()
    steps = [
        "Design Team Response: Review all findings and provide written response within 2 weeks.",
        "Variance Requests: Submit formal variance documentation for all items in Section 7.",
        "Resubmission: Updated drawings shall be submitted for the next phase review.",
        "Follow-up Review: WSU Facilities Services will conduct a follow-up review.",
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Number")

    # Sign-off table
    doc.add_heading("Review Sign-off", level=2)
    t = doc.add_table(rows=4, cols=4)
    t.style = "Table Grid"
    for c, h in enumerate(["Role", "Name", "Date", "Signature"]):
        t.rows[0].cells[c].text = h
    for i, role in enumerate(["Reviewer", "WSU Project Manager", "Design Team Lead"], 1):
        t.rows[i].cells[0].text = role
    format_table(t)

    # Appendix
    doc.add_page_break()
    doc.add_heading("Appendix A: Complete Checklist", level=1)
    add_crimson_rule()
    doc.add_paragraph("See: checklist.txt")
    p = doc.add_paragraph()
    run = p.add_run(
        "Non-compliant items in the checklist reference their Finding number (F-###) "
        "for cross-document traceability."
    )
    run.font.italic = True

    doc.save(os.path.join(output_dir, "report.docx"))

# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    if len(sys.argv) < 2:
        print("Usage: python generate_reports.py <review-data.json>")
        sys.exit(1)

    json_path = sys.argv[1]
    output_dir = os.path.dirname(os.path.abspath(json_path))

    print(f"Reading: {json_path}")
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # Override executive_summary from executive-summary.txt if present
    summary_file = os.path.join(output_dir, "executive-summary.txt")
    if os.path.isfile(summary_file):
        with open(summary_file, "r", encoding="utf-8") as f:
            text = f.read().strip()
            if text and text != "PLACEHOLDER":
                data.setdefault("narratives", {})["executive_summary"] = text
                print(f"  Using executive summary from {summary_file}")

    # Sort findings by severity then division
    findings = sorted(data.get("findings", []), key=sev_sort_key)

    # Validate
    print("Validating numbers...")
    errors = validate(data, findings)
    if errors:
        err_msg = "NUMBER VALIDATION FAILED:\n" + "\n".join(f"  - {e}" for e in errors)
        print(err_msg)
        with open(os.path.join(output_dir, "FAILED"), "w") as fh:
            fh.write(err_msg)
        sys.exit(1)
    print(f"  Validation passed: {len(findings)} findings, numbers balanced.")

    # Generate all outputs
    try:
        print("Generating findings.txt...")
        generate_findings_txt(findings, output_dir)

        print("Generating checklist.txt...")
        generate_checklist_txt(data, findings, output_dir)

        print("Generating notes.txt...")
        generate_notes_txt(data, output_dir)

        print("Generating report.xlsx...")
        generate_xlsx(data, findings, output_dir)

        print("Generating report.docx...")
        generate_docx(data, findings, output_dir)

        # All outputs written — create sentinel
        with open(os.path.join(output_dir, "COMPLETE"), "w") as fh:
            fh.write("")
        print("All reports generated. COMPLETE.")

    except Exception as e:
        err_msg = f"REPORT GENERATION FAILED:\n{traceback.format_exc()}"
        print(err_msg)
        with open(os.path.join(output_dir, "FAILED"), "w") as fh:
            fh.write(err_msg)
        sys.exit(1)


if __name__ == "__main__":
    main()
